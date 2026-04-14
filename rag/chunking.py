"""
rag/chunking.py
───────────────
Document ingestion and text chunking utilities.

Reads a .docx file and splits its content into overlapping chunks
optimised for embedding and retrieval.

Includes content hashing for incremental ingestion — only changed
chunks need new embeddings.
"""

from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document


# ─── Text Extraction ─────────────────────────────────────────────

def extract_text(docx_path: str | Path) -> str:
    """
    Extract all paragraph text from a Word (.docx) document.

    Also extracts text from tables to ensure no institute info is missed.

    Args:
        docx_path: Absolute or relative path to the .docx file.

    Returns:
        A single string with all paragraphs joined by newlines.

    Raises:
        FileNotFoundError: If the document does not exist.
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"Document not found: {docx_path}")

    doc = Document(str(docx_path))

    parts: list[str] = []

    # Paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)

    # Tables (many coaching-institute docs use tables for fee grids, etc.)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


# ─── Content Hashing ─────────────────────────────────────────────

def hash_chunk(text: str) -> str:
    """
    Generate a deterministic SHA-256 hash for a text chunk.

    Used by the incremental ingestion script to detect which chunks
    have changed and need re-embedding.
    """
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


# ─── Text Chunking ───────────────────────────────────────────────

def chunk_text(
    text: str,
    chunk_size: int = 500,
    chunk_overlap: int = 80,
) -> list[str]:
    """
    Split *text* into overlapping chunks of approximately *chunk_size*
    characters.  Overlap ensures context is not lost at chunk boundaries.

    The chunker tries to break on sentence boundaries to keep
    chunks semantically coherent.

    Args:
        text:          The full document text.
        chunk_size:    Target size of each chunk in characters.
        chunk_overlap: Number of overlapping characters between
                       consecutive chunks.

    Returns:
        A list of text chunks.
    """
    if not text:
        return []

    chunks: list[str] = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = start + chunk_size

        # Try to find a clean sentence boundary within the window
        if end < text_len:
            # Prefer period-space, then newline, then any period
            for sep in [". ", ".\n", "\n"]:
                boundary = text.rfind(sep, start, end)
                if boundary != -1 and boundary > start:
                    end = boundary + 1
                    break

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        # Ensure start always advances (prevent infinite loop)
        next_start = end - chunk_overlap if end < text_len else text_len
        if next_start <= start:
            next_start = start + 1
        start = next_start

    return chunks


def chunk_text_with_hashes(
    text: str,
    chunk_size: int = 500,
    chunk_overlap: int = 80,
) -> list[tuple[str, str]]:
    """
    Chunk text and return (content, content_hash) tuples.

    Used by the incremental ingestion script to diff against existing
    database rows.
    """
    chunks = chunk_text(text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return [(c, hash_chunk(c)) for c in chunks]
